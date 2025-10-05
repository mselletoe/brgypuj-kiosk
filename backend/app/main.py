from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import psycopg2
import os
from docx import Document
from io import BytesIO

# Database URL
DB_URL = os.getenv("DATABASE_URL", "postgresql://admin:admin7890@localhost:5432/kioskdb")

# Initialize FastAPI
app = FastAPI(title="Kiosk Backend API")

# Enable CORS for frontend requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Use your frontend URL in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Database connection
def get_conn():
    try:
        conn = psycopg2.connect(DB_URL)
        return conn
    except Exception as e:
        print("Database connection failed:", e)
        raise

@app.get("/")
def root():
    return {"message": "Kiosk backend is running"}

# Get paginated residents
@app.get("/api/residents")
def list_residents(page: int = 1, limit: int = 10):
    offset = (page - 1) * limit
    with get_conn() as conn:
        with conn.cursor() as cur:
            # Fetch paginated residents
            cur.execute(
                """
                SELECT id, last_name, first_name, middle_name, suffix, sex_gender, age,
                    birthdate, years_residency, unit_blk_street, purok
                FROM residents
                ORDER BY id
                LIMIT %s OFFSET %s
                """,
                (limit, offset)
            )
            rows = cur.fetchall()

            # Fetch total count for pagination
            cur.execute("SELECT COUNT(*) FROM residents")
            total = cur.fetchone()[0]

            data = [
                {
                    "id": r[0],
                    "last_name": r[1],
                    "first_name": r[2],
                    "middle_name": r[3],
                    "suffix": r[4],
                    "sex_gender": r[5],
                    "age": r[6],
                    "birthdate": r[7].isoformat() if r[7] else None,
                    "years_residency": r[8],
                    "unit_blk_street": r[9],
                    "purok": r[10],
                }
                for r in rows
            ]

            return {"data": data, "total": total}

# # Certificate request model
# class CertificateRequest(BaseModel):
#     resident_id: int

# # Generate Word certificate
# @app.post("/api/generate-certificate")
# def generate_certificate(req: CertificateRequest):
    with get_conn() as conn:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT first_name, last_name, brgy FROM residents WHERE id = %s",
                (req.resident_id,)
            )
            row = cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="Resident not found")
            first, last, brgy = row

    # Create Word document
    doc = Document()
    doc.add_heading("Barangay Certificate", level=1)
    doc.add_paragraph(f"This certifies that {first} {last} is a resident of {brgy}.")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return FileResponse(
        path_or_bytes=bio.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"certificate_{req.resident_id}.docx"
    )

# Run uvicorn if executed directly
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
